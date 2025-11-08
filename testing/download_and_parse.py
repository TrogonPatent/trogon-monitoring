"""
USPTO Bulk Patent Download & Parser
Downloads one week of patents and creates test files

Requirements:
pip install requests lxml python-docx reportlab

Usage:
python download_and_parse.py --week 241105 --count 2
python download_and_parse.py --week 241105 --count 1000
"""

import os
import sys
import json
import random
import zipfile
import argparse
import requests
from pathlib import Path
from lxml import etree
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

class USPTOParser:
    def __init__(self, output_dir="test_patents"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Create subdirectories
        self.download_dir = self.output_dir / "downloaded"
        self.specs_dir = self.output_dir / "specifications"
        self.drawings_dir = self.output_dir / "drawings"
        self.ground_truth_dir = self.output_dir / "ground_truth"
        
        for d in [self.download_dir, self.specs_dir, self.drawings_dir, self.ground_truth_dir]:
            d.mkdir(exist_ok=True)
        
        print(f"✓ Output directory: {self.output_dir}")
    
    def download_week(self, week):
        """
        Download one week of patents
        week format: YYMMDD (e.g., 241105 for Nov 5, 2024)
        """
        year = "20" + week[:2]
        filename = f"ipg{week}.zip"
        url = f"https://bulkdata.uspto.gov/data/patent/grant/redbook/fulltext/{year}/{filename}"
        
        zip_path = self.download_dir / filename
        
        if zip_path.exists():
            print(f"✓ Already downloaded: {filename}")
            return zip_path
        
        print(f"Downloading {url}...")
        print("This may take a few minutes (50-100 MB)...")
        
        try:
            response = requests.get(url, stream=True, timeout=300)
            response.raise_for_status()
            
            total_size = int(response.headers.get('content-length', 0))
            downloaded = 0
            
            with open(zip_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
                    downloaded += len(chunk)
                    if total_size > 0:
                        pct = (downloaded / total_size) * 100
                        print(f"\rProgress: {pct:.1f}%", end='', flush=True)
            
            print(f"\n✓ Downloaded: {filename} ({total_size / 1024 / 1024:.1f} MB)")
            return zip_path
            
        except Exception as e:
            print(f"\n✗ Download failed: {e}")
            if zip_path.exists():
                zip_path.unlink()
            sys.exit(1)
    
    def parse_patents(self, zip_path, max_count=None):
        """Extract and parse patents from ZIP file"""
        patents = []
        
        print(f"\nExtracting patents from {zip_path.name}...")
        
        with zipfile.ZipFile(zip_path, 'r') as zf:
            xml_files = [f for f in zf.namelist() if f.endswith('.xml')]
            
            if not xml_files:
                print("✗ No XML file found in ZIP")
                sys.exit(1)
            
            xml_file = xml_files[0]
            print(f"Parsing {xml_file}...")
            
            with zf.open(xml_file) as xf:
                # Read file content
                content = xf.read()
                
                # Check if this is 2025 format (multiple XML declarations)
                if content.count(b'<?xml') > 1:
                    print("Detected 2025 format (multiple XML declarations)")
                    patents = self.parse_2025_format(content, max_count)
                else:
                    print("Detected 2024 format (single XML document)")
                    patents = self.parse_2024_format(content, max_count)
        
        print(f"\n✓ Parsed {len(patents)} valid patents")
        return patents
    
    def parse_2024_format(self, content, max_count=None):
        """Parse 2024 format (single XML document)"""
        patents = []
        root = etree.fromstring(content)
        
        for elem in root.findall('.//us-patent-grant'):
            try:
                patent = self.parse_patent_xml(elem)
                if patent:
                    patents.append(patent)
                    
                    if len(patents) % 100 == 0:
                        print(f"\rParsed {len(patents)} patents...", end='', flush=True)
                    
                    if max_count and len(patents) >= max_count:
                        break
            except Exception as e:
                pass
        
        return patents
    
    def parse_2025_format(self, content, max_count=None):
        """Parse 2025 format (multiple XML declarations)"""
        patents = []
        
        # Split by XML declarations
        parts = content.split(b'<?xml version="1.0" encoding="UTF-8"?>')
        
        print(f"Found {len(parts)-1} patent documents")
        
        # Debug: Print structure of first patent
        first_patent_xml = b'<?xml version="1.0" encoding="UTF-8"?>' + parts[1]
        first_root = etree.fromstring(first_patent_xml)
        print("\n=== FIRST PATENT XML STRUCTURE ===")
        print(f"Root tag: {first_root.tag}")
        for child in first_root[:5]:  # First 5 children
            print(f"  Child: {child.tag}")
            for subchild in child[:3]:  # First 3 grandchildren
                print(f"    Grandchild: {subchild.tag}")
        
        # Check for CPC in different locations
        cpc_locations = [
            './/classification-cpc',
            './/classifications-cpc',
            './/cpc-classification',
            './/main-classification',
            './/classification-cpc-text'
        ]
        
        print("\n=== SEARCHING FOR CPC CODES ===")
        for loc in cpc_locations:
            found = first_root.findall(loc)
            if found:
                print(f"✓ Found {len(found)} elements at: {loc}")
                if found[0].text:
                    print(f"  Text: {found[0].text}")
            else:
                print(f"✗ Not found: {loc}")
        
        print("=" * 50)
        sys.exit(0)  # Stop here to review structure
        
        for i, part in enumerate(parts[1:], 1):  # Skip first empty part
            if not part.strip():
                continue
            
            try:
                # Add XML declaration back
                xml_doc = b'<?xml version="1.0" encoding="UTF-8"?>' + part
                
                # Parse this individual patent
                root = etree.fromstring(xml_doc)
                
                # In 2025 format, root IS the us-patent-grant element
                # Check if root tag is correct
                if 'us-patent-grant' not in root.tag:
                    continue
                
                patent = self.parse_patent_xml(root)
                if patent:
                    patents.append(patent)
                    
                    if len(patents) % 100 == 0:
                        print(f"\rParsed {len(patents)} patents...", end='', flush=True)
                    
                    if max_count and len(patents) >= max_count:
                        break
            
            except Exception as e:
                # Skip malformed patents
                if len(patents) == 0 and i < 5:  # Debug first few failures
                    print(f"\nDebug error on patent {i}: {e}")
        
        return patents
    
    def parse_patent_xml(self, elem):
        """Parse single patent XML element"""
        patent = {}
        
        # Patent number
        doc_num = elem.find('.//doc-number')
        if doc_num is None or not doc_num.text:
            return None
        patent['number'] = f"US{doc_num.text.strip()}"
        
        # Title
        title = elem.find('.//invention-title')
        if title is None or not title.text:
            print(f"  Skipping {patent['number']}: No title")
            return None
        patent['title'] = title.text.strip()
        
        # Publication date
        pub_date = elem.find('.//publication-reference/document-id/date')
        patent['publication_date'] = pub_date.text if pub_date is not None else "Unknown"
        
        # Assignee
        assignee = elem.find('.//assignee/orgname')
        patent['assignee'] = assignee.text.strip() if assignee is not None and assignee.text else "Individual"
        
        # CPC Classifications (ground truth)
        cpc_codes = []
        
        # Primary CPC
        main_cpc = elem.find('.//classification-cpc/main-cpc/classification-cpc-text')
        if main_cpc is not None and main_cpc.text:
            patent['primary_cpc'] = main_cpc.text.strip()
            cpc_codes.append({
                'code': main_cpc.text.strip(),
                'type': 'primary'
            })
        else:
            print(f"  Skipping {patent['number']}: No primary CPC")
            return None  # Skip if no primary CPC
        
        # Additional CPCs
        for further_cpc in elem.findall('.//classification-cpc/further-cpc/classification-cpc-text'):
            if further_cpc.text:
                cpc_codes.append({
                    'code': further_cpc.text.strip(),
                    'type': 'secondary'
                })
        
        patent['all_cpc_codes'] = cpc_codes
        
        # Abstract
        abstract_parts = []
        for p in elem.findall('.//abstract/p'):
            text = ''.join(p.itertext()).strip()
            if text:
                abstract_parts.append(text)
        patent['abstract'] = '\n\n'.join(abstract_parts) if abstract_parts else ""
        
        # Description (specification WITHOUT claims)
        desc_parts = []
        description = elem.find('.//description')
        if description is not None:
            for p in description.findall('.//p'):
                text = ''.join(p.itertext()).strip()
                if text:
                    desc_parts.append(text)
        
        patent['description'] = '\n\n'.join(desc_parts) if desc_parts else ""
        
        # Claims (ground truth only)
        claims = []
        for claim in elem.findall('.//claims/claim'):
            claim_id = claim.get('id', '')
            claim_text_elem = claim.find('.//claim-text')
            if claim_text_elem is not None:
                claim_text = ''.join(claim_text_elem.itertext()).strip()
                if claim_text:
                    claims.append({
                        'id': claim_id,
                        'text': claim_text
                    })
        patent['claims'] = claims
        
        # Skip if missing critical data
        if not patent['description'] or len(patent['description']) < 500:
            print(f"  Skipping {patent['number']}: Description too short ({len(patent['description'])} chars)")
            return None  # Too short
        if not patent['claims']:
            print(f"  Skipping {patent['number']}: No claims")
            return None  # No claims
        
        print(f"  ✓ Valid patent: {patent['number']}")
        return patent
    
    def create_spec_docx(self, patent):
        """Create DOCX specification (NO claims, NO CPC codes)"""
        doc = Document()
        
        # Title
        title_para = doc.add_heading(patent['title'], level=1)
        title_para.alignment = 1  # Center
        
        # Metadata
        doc.add_paragraph(f"Patent Number: {patent['number']}")
        doc.add_paragraph(f"Publication Date: {patent['publication_date']}")
        doc.add_paragraph(f"Assignee: {patent['assignee']}")
        doc.add_paragraph()
        
        # Abstract
        if patent['abstract']:
            doc.add_heading('ABSTRACT', level=2)
            doc.add_paragraph(patent['abstract'])
            doc.add_paragraph()
        
        # Description
        doc.add_heading('DETAILED DESCRIPTION', level=2)
        
        # Split description into paragraphs
        for para in patent['description'].split('\n\n'):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.style.font.size = Pt(11)
        
        # Save
        filename = f"{patent['number']}-spec.docx"
        filepath = self.specs_dir / filename
        doc.save(filepath)
        
        return filepath
    
    def create_drawing_pdf(self, patent):
        """Create placeholder PDF for drawing"""
        filename = f"{patent['number']}-drawing.pdf"
        filepath = self.drawings_dir / filename
        
        c = canvas.Canvas(str(filepath), pagesize=letter)
        
        # Title
        c.setFont("Helvetica-Bold", 16)
        c.drawString(100, 700, f"Patent {patent['number']}")
        
        # Placeholder text
        c.setFont("Helvetica", 12)
        c.drawString(100, 670, "Drawing Placeholder")
        c.drawString(100, 650, "(Actual drawings not included in bulk XML data)")
        
        c.save()
        
        return filepath
    
    def create_ground_truth(self, patent):
        """Create ground truth JSON"""
        ground_truth = {
            'patent_number': patent['number'],
            'title': patent['title'],
            'publication_date': patent['publication_date'],
            'assignee': patent['assignee'],
            
            'ground_truth_cpc': {
                'primary': patent['primary_cpc'],
                'all_codes': patent['all_cpc_codes']
            },
            
            'ground_truth_claims': {
                'independent_claims': [
                    c['text'] for c in patent['claims'] 
                    if c['id'] in ['1', '10', '20', '30']
                ],
                'all_claims': [c['text'] for c in patent['claims']],
                'claim_count': len(patent['claims'])
            },
            
            'specification_text_length': len(patent['description']),
            'abstract': patent['abstract']
        }
        
        filename = f"{patent['number']}-ground-truth.json"
        filepath = self.ground_truth_dir / filename
        
        with open(filepath, 'w') as f:
            json.dump(ground_truth, f, indent=2)
        
        return filepath
    
    def process_patents(self, patents):
        """Create all test files for patents"""
        print(f"\nCreating test files for {len(patents)} patents...")
        
        processed = []
        errors = []
        
        for i, patent in enumerate(patents, 1):
            try:
                spec_path = self.create_spec_docx(patent)
                drawing_path = self.create_drawing_pdf(patent)
                gt_path = self.create_ground_truth(patent)
                
                processed.append({
                    'patent_number': patent['number'],
                    'spec_path': str(spec_path),
                    'drawing_path': str(drawing_path),
                    'ground_truth_path': str(gt_path)
                })
                
                if i % 50 == 0:
                    print(f"✓ Processed {i}/{len(patents)} patents")
            
            except Exception as e:
                errors.append({
                    'patent': patent['number'],
                    'error': str(e)
                })
        
        print(f"\n✓ Successfully processed: {len(processed)}")
        if errors:
            print(f"✗ Errors: {len(errors)}")
        
        # Save processing manifest
        manifest = {
            'processed': processed,
            'errors': errors,
            'total': len(patents)
        }
        
        manifest_path = self.output_dir / 'processing_manifest.json'
        with open(manifest_path, 'w') as f:
            json.dump(manifest, f, indent=2)
        
        print(f"✓ Manifest saved: {manifest_path}")
        
        return processed

def main():
    parser = argparse.ArgumentParser(description='Download and parse USPTO patents')
    parser.add_argument('--week', help='Week to download (YYMMDD format, e.g., 241105)')
    parser.add_argument('--zip-file', help='Path to existing ZIP file (skip download)')
    parser.add_argument('--count', type=int, default=2, help='Number of patents to process (default: 2 for testing)')
    parser.add_argument('--output', default='test_patents', help='Output directory (default: test_patents)')
    
    args = parser.parse_args()
    
    if not args.week and not args.zip_file:
        print("Error: Must specify either --week or --zip-file")
        sys.exit(1)
    
    print("=" * 60)
    print("USPTO Bulk Patent Download & Parser")
    print("=" * 60)
    print(f"Count: {args.count}")
    print(f"Output: {args.output}")
    print("=" * 60)
    
    # Initialize parser
    parser = USPTOParser(output_dir=args.output)
    
    # Get ZIP file (download or use existing)
    if args.zip_file:
        zip_path = Path(args.zip_file)
        if not zip_path.exists():
            print(f"✗ File not found: {args.zip_file}")
            sys.exit(1)
        print(f"Using existing ZIP: {zip_path}")
    else:
        zip_path = parser.download_week(args.week)
    
    # Parse patents (will stop after reaching count)
    patents = parser.parse_patents(zip_path, max_count=args.count * 2)  # Parse extra in case some are filtered
    
    # Randomly sample if we have more than needed
    if len(patents) > args.count:
        patents = random.sample(patents, args.count)
        print(f"✓ Randomly sampled {args.count} patents")
    
    # Create test files
    processed = parser.process_patents(patents)
    
    print("\n" + "=" * 60)
    print("COMPLETE!")
    print("=" * 60)
    print(f"Files created in: {parser.output_dir}")
    print(f"  Specifications: {len(processed)}")
    print(f"  Drawings: {len(processed)}")
    print(f"  Ground truth: {len(processed)}")
    print("\nNext step: Run upload_test.py to test API integration")

if __name__ == '__main__':
    main()
